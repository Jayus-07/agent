"""评测集 Ground Truth 全量校验脚本。

背景：2026-08-19 发现 12 条用例的 expected doc_id/snippet 与 KB 实际内容不符
（commit 7f32825 只修了 5 条）。本脚本对评测集做全量审计，防止错标再次溜进基线。

校验项：
  1. 结构：id 唯一、question 无逐字重复、必填字段齐全
  2. doc_id 协议：relevant_docs 必须是 md5(basename)[:10] 且对应 KB 目录中真实存在的文件 [ERROR]
  3. snippet 真实性：
     - match_type=snippet 用例（snippet 是判据）：关键词必须存在于 KB 任意文档 [ERROR]
     - doc-bound 用例（snippet 仅展示，doc_id 才是判据）：关键词未出现在期望文档
       记 [WARN]（无协议风险，但提示标注与原文措辞不一致）
  4. 负样本一致性：should_reject=True 的用例不得标注 relevant_docs/relevant_snippets [ERROR]

用法：
    python backend/scripts/validate_eval_dataset.py
    python backend/scripts/validate_eval_dataset.py --dataset rag_test_kb.json
    python backend/scripts/validate_eval_dataset.py --kb-dir data/docs/rag_test_kb/general

退出码：0 = 全部通过；1 = 存在 ERROR（WARN 不阻断）。
"""
import argparse
import hashlib
import json
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATASET_DIR = PROJECT_ROOT / "backend" / "evaluation" / "datasets"
DEFAULT_KB_DIR = PROJECT_ROOT / "data" / "docs" / "rag_test_kb" / "general"


def normalize(text: str) -> str:
    """与 runner _match_by_snippet 一致的归一化：全角转半角 + 去全部空白。"""
    text = text.translate(
        {i: i - 0xFEE0 for i in range(0xFF01, 0xFF5F)}
    ).replace("\u3000", " ")
    return "".join(text.split())


def doc_id_of(filename: str) -> str:
    """doc_id 协议：md5(basename)[:10]。"""
    return hashlib.md5(filename.encode("utf-8")).hexdigest()[:10]


def extract_text(path: Path) -> str | None:
    """按扩展名解析文档文本；解析失败返回 None（记 WARN 而非 ERROR）。"""
    try:
        if path.suffix.lower() == ".md":
            return path.read_text(encoding="utf-8", errors="replace")
        if path.suffix.lower() == ".docx":
            import docx
            doc = docx.Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                parts.extend(
                    cell.text for row in table.rows for cell in row.cells
                )
            return "\n".join(parts)
        if path.suffix.lower() == ".pdf":
            import fitz
            with fitz.open(str(path)) as pdf:
                return "\n".join(page.get_text() for page in pdf)
    except Exception as e:  # noqa: BLE001
        print(f"    [WARN] 解析失败 {path.name}: {e}")
    return None


def main() -> int:
    parser = argparse.ArgumentParser(description="评测集 ground truth 全量校验")
    parser.add_argument("--dataset", default="rag_test_kb.json",
                        help="评测集文件名（相对 datasets/ 目录）")
    parser.add_argument("--kb-dir", default=str(DEFAULT_KB_DIR),
                        help="KB 源文档目录")
    args = parser.parse_args()

    dataset_path = DATASET_DIR / args.dataset
    kb_dir = Path(args.kb_dir)
    if not dataset_path.exists():
        print(f"[ERROR] 评测集不存在: {dataset_path}")
        return 1
    if not kb_dir.is_dir():
        print(f"[ERROR] KB 目录不存在: {kb_dir}")
        return 1

    with open(dataset_path, encoding="utf-8") as f:
        data = json.load(f)
    cases = data["test_cases"]

    # KB 文件 → doc_id 映射（递归扫描，支持 --kb-dir 传含子目录的 KB 根目录）
    kb_files: dict[str, Path] = {}
    for p in kb_dir.rglob("*"):
        if p.is_file() and not p.name.startswith("."):
            kb_files[doc_id_of(p.name)] = p
    print(f"评测集: {dataset_path.name} ({len(cases)} 条) | "
          f"KB 文档: {len(kb_files)} 篇 | version={data.get('version', '?')}")

    errors: list[str] = []
    warns: list[str] = []
    seen_ids: set[str] = set()
    seen_questions: dict[str, str] = {}
    text_cache: dict[str, str | None] = {}

    for case in cases:
        cid = case.get("id", "<no-id>")
        question = case.get("question", "")
        expected = case.get("expected", {})
        relevant_docs = expected.get("relevant_docs") or []
        snippets = expected.get("relevant_snippets") or []
        should_reject = expected.get("should_reject", False)

        # 1. 结构校验
        if cid in seen_ids:
            errors.append(f"{cid}: id 重复")
        seen_ids.add(cid)
        nq = normalize(question)
        if nq in seen_questions:
            errors.append(f"{cid}: question 与 {seen_questions[nq]} 逐字重复: {question}")
        seen_questions[nq] = cid
        if not case.get("kb_id"):
            errors.append(f"{cid}: 缺少 kb_id 字段")
        difficulty = case.get("metadata", {}).get("difficulty")
        if difficulty not in ("easy", "medium", "hard"):
            warns.append(f"{cid}: difficulty 非标准值: {difficulty}")

        # 4. 负样本一致性
        if should_reject:
            if relevant_docs or snippets:
                errors.append(f"{cid}: should_reject=True 但标注了 relevant_docs/snippets")
            continue

        # snippet-only 用例（match_type=snippet，不绑定 doc_id）：
        # snippet 必须存在于 KB 任意文档中
        match_type = expected.get("match_type")
        if match_type == "snippet":
            if not snippets:
                errors.append(f"{cid}: match_type=snippet 但未标注 relevant_snippets")
                continue
            merged_all = ""
            for p in kb_files.values():
                key = str(p)
                if key not in text_cache:
                    text_cache[key] = extract_text(p)
                if text_cache[key] is not None:
                    merged_all += "\n" + text_cache[key]
            if not merged_all:
                warns.append(f"{cid}: KB 文档均无法解析，跳过 snippet 校验")
                continue
            normalized_all = normalize(merged_all)
            for s in snippets:
                if normalize(s) not in normalized_all:
                    errors.append(
                        f"{cid}: snippet '{s}' 未出现在 KB 任何文档中（ground truth 错标）"
                    )
            continue

        # 2. doc_id 协议校验
        if not relevant_docs:
            errors.append(f"{cid}: 正样本缺少 relevant_docs（且未标 should_reject）")
            continue
        doc_paths: list[Path] = []
        for did in relevant_docs:
            if did in kb_files:
                doc_paths.append(kb_files[did])
            else:
                errors.append(
                    f"{cid}: doc_id '{did}' 在 KB 目录中无对应文件"
                    f"（协议 md5(basename)[:10]，检查文件是否改名/未入库）"
                )

        # 3. snippet 真实性校验（期望文档文本并集）
        if snippets and doc_paths:
            merged = ""
            for p in doc_paths:
                key = str(p)
                if key not in text_cache:
                    text_cache[key] = extract_text(p)
                if text_cache[key] is not None:
                    merged += "\n" + text_cache[key]
            if not merged:
                warns.append(f"{cid}: 期望文档均无法解析，跳过 snippet 校验")
                continue
            normalized_merged = normalize(merged)
            for s in snippets:
                if normalize(s) not in normalized_merged:
                    # doc-bound 用例中 snippet 仅展示（判据是 doc_id），降级为 WARN；
                    # 提示标注关键词与文档原文措辞不一致，建议向原文对齐
                    warns.append(
                        f"{cid}: snippet '{s}' 未出现在期望文档 "
                        f"{[p.name for p in doc_paths]} 中（展示字段，不影响判分，"
                        f"建议改为文档原文措辞）"
                    )

    # 汇总
    for w in warns:
        print(f"  [WARN] {w}")
    for e in errors:
        print(f"  [ERROR] {e}")
    if errors:
        print(f"\n[FAIL] 校验失败：{len(errors)} 个错误，{len(warns)} 个警告")
        return 1
    print(f"\n[OK] 校验通过：{len(cases)} 条用例全部合规（{len(warns)} 个警告）")
    return 0


if __name__ == "__main__":
    # Windows GBK 终端下避免中文/emoji 编码崩溃
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.exit(main())
