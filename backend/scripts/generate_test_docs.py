"""生成 11 个 rag_test_kb 测试文件 — 内容虚构但结构贴近业务。

按文件名语义,每种类型挑不同的 chunking 路径来测:
  - .md / .docx → markdown parser + chunking_router(可能走结构/QA/Recursive)
  - .pdf         → pdf parser + 各种 chunk strategy

设计目标:每个文件覆盖至少一种 chunking 策略路径,便于后续回归测试。

所有文件:仅用于测试,内容虚构,不构成真实业务建议。
"""
from __future__ import annotations

import os
import random
import sys
from pathlib import Path

sys.path.insert(0, '.')

# 让 _reportlab_chinese_helper 可被同目录模块导入
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _reportlab_chinese_helper import apply_chinese_styles  # noqa: E402

OUT_DIR = Path(r"D:\Program Files\workplace\agent\data\docs\rag_test_kb\general")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# 一些虚构/看起来真实的字段(测试用,不要当真)
COMPANY = "示例跨境电商公司"
DISCLAIMER = "【仅供测试 · 内容虚构】"

# 报告lab 中文支持:注册 Windows 系统字体(必须,否则 PDF 抽取后中文全乱码 → filter 全部 reject)
_REPORTLAB_FONT_REGISTERED = False


def _ensure_reportlab_chinese_font():
    """注册 reportlab 支持中文的字体(优先用 CID 字体,内置 ToUnicode CMap)。

    背景:reportlab 默认 Helvetica 没中文 cmap。
    关键:PyMuPDF 提取中文需要 PDF 内嵌 ToUnicode CMap,所以优先用
    reportlab 自带的 CID 字体(STSong-Light),而不是 TrueType 字体。
    TrueType 字体(simhei 等)在 reportlab 不会自动生成 CMap,PyMuPDF
    提取出来仍是乱码(虽然不是 ? 了)。

    字体选择:
      1. STSong-Light (reportlab 内置 CID,自带 CMap,首选)
      2. 系统 TTFont (simhei / msyh 等,但 PDF 提取仍有问题,兜底)
    """
    global _REPORTLAB_FONT_REGISTERED
    if _REPORTLAB_FONT_REGISTERED:
        return "STSong-Light"
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        # 首选:STSong-Light (CID 字体,内置 ToUnicode)
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        _REPORTLAB_FONT_REGISTERED = True
        return "STSong-Light"
    except Exception as e:
        # 兜底:TTFont(不一定保证中文 PDF 文本可提取)
        try:
            from reportlab.pdfbase.ttfonts import TTFont
            for path in [
                r"C:\Windows\Fonts\simhei.ttf",
                r"C:\Windows\Fonts\msyh.ttc",
                "/System/Library/Fonts/PingFang.ttc",
            ]:
                if os.path.exists(path):
                    pdfmetrics.registerFont(TTFont("simhei", path))
                    _REPORTLAB_FONT_REGISTERED = True
                    return "simhei"
        except Exception:
            pass
        return "Helvetica"


CHINESE_FONT = _ensure_reportlab_chinese_font()


# ============ .md / .docx ============

def gen_01_faq_md() -> None:
    """FAQ — 走 QAChunkStrategy 的种子数据:清晰 Q/A 配对。"""
    qas = [
        ("客户下单后多久发货?", "现货商品 24 小时内发出,预售商品以商品详情页标注的发货时间为准。"),
        ("支持哪些支付方式?", "支持支付宝、微信支付、银联及境外信用卡(Visa/Mastercard)。"),
        ("如何申请退款?", "在「我的订单」点击「申请退款」,填写原因后提交,客服 1-2 个工作日内审核。"),
        ("退货运费谁承担?", "质量问题由商家承担;无理由退货由买家承担(商品详情页另有约定的除外)。"),
        ("订单可以修改地址吗?", "待发货状态可自行修改;已发货状态请联系客服申请拦截改址。"),
        ("发票如何开具?", "在订单详情页提交开票申请,电子发票 3 个工作日内发至预留邮箱。"),
        ("会员积分如何累计?", "每消费 1 元累计 1 分,满 1000 分可抵扣 10 元。"),
        ("海外配送时效?", "标准配送 7-15 个工作日,加急配送 3-7 个工作日,具体以物流追踪页为准。"),
        ("商品保修期多久?", "电子产品按品牌政策,一般 1 年;其他品类详见商品详情页。"),
        ("如何联系人工客服?", "在线客服 9:00-22:00,电话客服 400-000-0000 转 1。"),
        ("优惠券能否叠加使用?", "店铺券与平台券可叠加,但单笔订单最多使用 3 张优惠券。"),
        ("退货商品需要原包装吗?", "建议保持原包装;若包装损坏可能影响退款金额。"),
        ("如何查询物流轨迹?", "「我的订单」→「查看物流」可看到揽收、运输、派送、签收全节点。"),
        ("地址写错了能改吗?", "待发货状态可修改一次,已发货请通过客服申请。"),
        ("能开发票抬头吗?", "支持个人/企业抬头,需提供税号及开户银行(企业)。"),
    ]
    lines = [f"# {COMPANY} 客服 FAQ 知识库", DISCLAIMER, ""]
    for q, a in qas:
        lines.append(f"Q：{q}")
        lines.append(f"A：{a}")
        lines.append("")
    (OUT_DIR / "01_FAQ.md").write_text("\n".join(lines), encoding="utf-8")


def gen_11_purchase_contract_md() -> None:
    """采购合同 — 走 StructureChunkStrategy。包含分级标题 + 条款。"""
    body = f"""# 跨境电商采购合同

{DISCLAIMER}

## 一、合同双方

### 1.1 甲方(采购方)
- 名称:{COMPANY}
- 地址:示例市示例区示例路 100 号
- 法定代表人:张三

### 1.2 乙方(供应方)
- 名称:示例供应商有限公司
- 地址:示例市供应商园区 B 座
- 法定代表人:李四

## 二、标的物

| 商品名称 | 规格 | 数量(件) | 单价(元) | 金额(元) |
|---------|------|---------|----------|----------|
| 跨境商品 A | 标准 | 1000 | 50 | 50000 |
| 跨境商品 B | 加厚 | 500 | 80 | 40000 |

合计:人民币玖万元整(¥90,000.00)。

## 三、交付条款

3.1 交付时间:合同签订后 30 个自然日内完成全部交付。
3.2 交付地点:甲方指定仓库。
3.3 运输方式:乙方负责运输并承担运费。

## 四、质量要求

4.1 商品质量须符合国家相关标准及甲方书面要求。
4.2 抽检不合格率超过 5% 时,甲方有权要求换货或退货。

## 五、付款方式

5.1 合同签订后 3 个工作日内支付 30% 预付款。
5.2 货物验收合格后 7 个工作日内支付 65% 货款。
5.3 质保期满(交付后 90 日)无质量问题,支付剩余 5% 尾款。

## 六、违约责任

6.1 任何一方违约,应赔偿对方因此遭受的实际损失。
6.2 不可抗力按《民法典》相关规定处理。

## 七、争议解决

合同履行过程中发生的争议,双方协商解决;协商不成的,提交甲方所在地有管辖权的人民法院诉讼解决。

## 八、其他

合同自双方盖章之日起生效,一式两份,甲乙双方各执一份。
"""
    (OUT_DIR / "11_采购合同.md").write_text(body, encoding="utf-8")


# ============ .docx ============

def gen_02_after_sales_faq_docx() -> None:
    """售后 FAQ — docx 走 markdown 解析(QA 解析失败 → fallback Recursive)。"""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    title = doc.add_heading(f"{COMPANY} 售后服务 FAQ", level=0)
    doc.add_paragraph(DISCLAIMER)

    items = [
        ("七天无理由退货范围", "支持,客户在签收后 7 天内可申请无理由退货(贴身衣物/食品等除外)。"),
        ("退货商品运费", "质量问题由商家承担运费;无理由退货由买家承担运费。"),
        ("退款时效", "商品验收合格后,1-3 个工作日内原路退回支付账户。"),
        ("换货流程", "在订单详情页提交「换货申请」,客服审核通过后寄回,仓库验收后重新发出。"),
        ("维修申请", "电子类商品保修期内非人为损坏可申请免费维修;人为损坏按成本价收费。"),
        ("发票丢失补开", "电子发票可在订单详情页重新下载;纸质发票需联系客服走流程。"),
        ("会员等级升降", "累计消费满 5000 升金卡,享 95 折;满 20000 升钻石,享 9 折。"),
        ("签收异常处理", "快递显示签收但未收到,24 小时内联系客服并提供运单号核查。"),
    ]
    for q, a in items:
        p = doc.add_paragraph()
        run = p.add_run("Q:" + q)
        run.bold = True
        doc.add_paragraph("A:" + a)
        doc.add_paragraph("")
    doc.save(OUT_DIR / "02_售后FAQ.docx")


def gen_04_purchase_process_docx() -> None:
    """采购流程 — docx 结构清晰(标题+步骤列表),走 StructureChunkStrategy。"""
    from docx import Document
    doc = Document()
    doc.add_heading(f"{COMPANY} 采购流程", level=0)
    doc.add_paragraph(DISCLAIMER)

    doc.add_heading("一、适用范围", level=1)
    doc.add_paragraph("本流程适用于所有商品采购活动,涵盖国内采购与跨境采购。")

    doc.add_heading("二、采购申请", level=1)
    doc.add_paragraph("2.1 需求部门在 OA 提交采购申请,注明商品名称、规格、数量、预算及用途。")
    doc.add_paragraph("2.2 单笔金额 5000 元以下由部门负责人审批;5000-50000 元需总监审批;50000 元以上需总经理审批。")

    doc.add_heading("三、供应商选择", level=1)
    doc.add_paragraph("3.1 合格供应商名录内的供应商优先。")
    doc.add_paragraph("3.2 新供应商需经资质初审、样品验证、试单三个环节,合格后纳入名录。")

    doc.add_heading("四、合同签订", level=1)
    doc.add_paragraph("4.1 单笔金额 10000 元以上必须签订书面合同。")
    doc.add_paragraph("4.2 合同条款需经法务审核,包含价格、质量、交期、违约责任。")

    doc.add_heading("五、订单执行", level=1)
    doc.add_paragraph("5.1 采购员在 ERP 下单,确认供应商及交付时间。")
    doc.add_paragraph("5.2 大宗采购需分批下单,首批不低于总数 30%。")

    doc.add_heading("六、验收入库", level=1)
    doc.add_paragraph("6.1 仓库按 PO 核对数量与规格。")
    doc.add_paragraph("6.2 质检员按 AQL 抽样标准检验,合格批次办理入库。")

    doc.add_heading("七、对账付款", level=1)
    doc.add_paragraph("7.1 供应商按合同账期提交对账单。")
    doc.add_paragraph("7.2 财务核对后按合同约定方式付款。")

    doc.save(OUT_DIR / "04_采购流程.docx")


def gen_07_logistics_sop_docx() -> None:
    """物流 SOP — docx,按 SOP 结构(StepChunkStrategy)。"""
    from docx import Document
    doc = Document()
    doc.add_heading(f"{COMPANY} 物流 SOP 标准操作流程", level=0)
    doc.add_paragraph(DISCLAIMER)

    doc.add_heading("目的", level=1)
    doc.add_paragraph("规范跨境物流操作,降低丢件破损率,提升客户满意度。")

    doc.add_heading("适用范围", level=1)
    doc.add_paragraph("仓储-国内段、跨境头程、目的国尾程。")

    doc.add_heading("操作步骤", level=1)

    steps = [
        ("步骤 1:入库扫描", "商品入库时逐件扫描,记录重量与体积,绑定 SKU。"),
        ("步骤 2:分拣打包", "按目的国分拣,核对收件人信息,使用合格包材。"),
        ("步骤 3:头程交接", "批量交接至国际物流商,获取主运单号。"),
        ("步骤 4:跨境运输", "海运/空运/快船,实时跟踪在途状态。"),
        ("步骤 5:清关申报", "提供商业发票、装箱单,按目的国法规申报。"),
        ("步骤 6:尾程派送", "目的国本地物流完成末端派送,获取签收凭证。"),
        ("步骤 7:异常处理", "丢件 48 小时核查,破损 24 小时补发,清关异常 2 个工作日补材料。"),
        ("步骤 8:数据归档", "运单全程状态归档,保留 3 年供查询。"),
    ]
    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_heading("考核指标", level=1)
    doc.add_paragraph("丢件率 ≤ 0.5%,破损率 ≤ 0.3%,平均妥投时效 ≤ 12 个工作日。")

    doc.save(OUT_DIR / "07_物流SOP.docx")


# ============ .pdf ============

def gen_03_inventory_pdf() -> None:
    """库存管理制度 — PDF 走 StructureChunkStrategy(标题明确)。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )
    from reportlab.lib import colors

    doc = SimpleDocTemplate(str(OUT_DIR / "03_库存管理制度.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph(f"<b>{COMPANY} 库存管理制度</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>第一章 总则</b>", styles["Heading2"]))
    body.append(Paragraph("第一条 目的:规范库存管理,降低呆滞库存,提升资金周转率。", styles["BodyText"]))
    body.append(Paragraph("第二条 适用范围:成品仓、原料仓、在途仓、跨境保税仓。", styles["BodyText"]))

    body.append(Paragraph("<b>第二章 库存分类</b>", styles["Heading2"]))
    body.append(Paragraph("库存按状态分五类:", styles["BodyText"]))
    data = [
        ["分类", "定义", "示例"],
        ["可用库存", "可立即销售的成品", "仓库正常出库"],
        ["在途库存", "已发货未签收的货物", "跨境头程在途"],
        ["锁定库存", "已售未发货或预留", "待发货订单库存"],
        ["残次库存", "损坏或过期的商品", "退回检验不合格品"],
        ["滞销库存", "连续 90 天无动销的商品", "老款服装"],
    ]
    t = Table(data, colWidths=[60, 200, 200])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    body.append(t)
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>第三章 出入库流程</b>", styles["Heading2"]))
    body.append(Paragraph("入库流程:采购到货 → 验收 → 质检 → 上架。", styles["BodyText"]))
    body.append(Paragraph("出库流程:订单审核 → 拣货 → 复核 → 打包 → 交接物流。", styles["BodyText"]))

    body.append(Paragraph("<b>第四章 盘点制度</b>", styles["Heading2"]))
    body.append(Paragraph("月度盘点:每月末由财务组织抽盘,差异率不超过 1%。", styles["BodyText"]))
    body.append(Paragraph("年度盘点:每年 12 月 31 日全员盘点,差异率不超过 0.5%。", styles["BodyText"]))

    body.append(Paragraph("<b>第五章 滞销处理</b>", styles["Heading2"]))
    body.append(Paragraph("连续 90 天无动销的商品,转入滞销池,启动促销或调拨方案。", styles["BodyText"]))

    body.append(PageBreak())
    body.append(Paragraph("<b>附录 A:常用单据</b>", styles["Heading2"]))
    body.append(Paragraph("入库单、出库单、调拨单、盘点表、损益单。", styles["BodyText"]))
    doc.build(body)


def gen_05_product_spec_pdf() -> None:
    """商品规格说明 — PDF 表格驱动,走 StructureChunkStrategy。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(str(OUT_DIR / "05_商品规格说明.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph(f"<b>{COMPANY} 跨境商品规格说明</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>一、概述</b>", styles["Heading2"]))
    body.append(Paragraph("本文档描述跨境商品的规格、包装与合规要求,适用于所有在售 SKU。", styles["BodyText"]))

    body.append(Paragraph("<b>二、规格参数</b>", styles["Heading2"]))
    data = [
        ["SKU", "品类", "重量(g)", "尺寸(cm)", "电压(V)", "包装"],
        ["SKU-001", "电水壶", "850", "20×15×20", "220", "彩盒"],
        ["SKU-002", "电水壶", "920", "22×16×22", "110", "彩盒"],
        ["SKU-003", "搅拌机", "1200", "25×20×25", "220", "彩盒"],
        ["SKU-004", "空气炸锅", "3500", "35×30×30", "220", "纸箱"],
        ["SKU-005", "咖啡机", "4200", "40×30×35", "220", "纸箱"],
    ]
    t = Table(data, colWidths=[60, 80, 60, 110, 50, 60])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    body.append(t)
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>三、包装规范</b>", styles["Heading2"]))
    body.append(Paragraph("3.1 主包装:彩盒或纸箱,外贴 SKU 条码与中文标签。", styles["BodyText"]))
    body.append(Paragraph("3.2 缓冲材料:珍珠棉或气泡膜,厚度不低于 2cm。", styles["BodyText"]))
    body.append(Paragraph("3.3 跨境标签:CE/FCC/RoHS 标识按目的国要求加贴。", styles["BodyText"]))

    body.append(Paragraph("<b>四、合规要求</b>", styles["Heading2"]))
    body.append(Paragraph("欧盟 CE、美国 FCC、日本 PSE、英国 UKCA 视目的国清关需要。", styles["BodyText"]))
    body.append(Paragraph("危险品(D 类电池)需提供 MSDS 与 UN38.3 测试报告。", styles["BodyText"]))

    doc.build(body)


def gen_06_return_policy_pdf() -> None:
    """退货政策 — PDF 多章节,走 StructureChunkStrategy。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    doc = SimpleDocTemplate(str(OUT_DIR / "06_退货政策.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph(f"<b>{COMPANY} 退货政策</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>一、适用范围</b>", styles["Heading2"]))
    body.append(Paragraph("所有在 {COMPANY} 平台销售的商品,除非商品详情页另有说明。".format(COMPANY=COMPANY), styles["BodyText"]))

    body.append(Paragraph("<b>二、七天无理由退货</b>", styles["Heading2"]))
    body.append(Paragraph("2.1 适用条件:商品不影响二次销售,包装配件齐全。", styles["BodyText"]))
    body.append(Paragraph("2.2 时间窗口:客户签收后 7 天(含)内申请。", styles["BodyText"]))
    body.append(Paragraph("2.3 运费承担:买家承担来回运费。", styles["BodyText"]))

    body.append(Paragraph("<b>三、质量问题退货</b>", styles["Heading2"]))
    body.append(Paragraph("3.1 商品存在质量问题(破损/错发/性能故障),全额退款,商家承担运费。", styles["BodyText"]))
    body.append(Paragraph("3.2 客户需在签收后 48 小时内上传照片凭证。", styles["BodyText"]))

    body.append(Paragraph("<b>四、不可退货商品</b>", styles["Heading2"]))
    body.append(Paragraph("贴身衣物、食品、定制商品、已拆封的密封包装商品(质量问题除外)。", styles["BodyText"]))

    body.append(Paragraph("<b>五、退款时效</b>", styles["Heading2"]))
    body.append(Paragraph("仓库验收合格后 1-3 个工作日原路退回支付账户。", styles["BodyText"]))

    body.append(Paragraph("<b>六、争议处理</b>", styles["Heading2"]))
    body.append(Paragraph("客户与商家协商不成的,可申请平台介入,平台裁决 7 个工作日内作出。", styles["BodyText"]))

    doc.build(body)


def gen_08_reimbursement_pdf() -> None:
    """员工报销制度 — PDF 走 StructureChunkStrategy(纯文本,无表)。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    doc = SimpleDocTemplate(str(OUT_DIR / "08_员工报销制度.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph(f"<b>{COMPANY} 员工报销制度</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>第一章 总则</b>", styles["Heading2"]))
    body.append(Paragraph("1.1 目的:规范员工费用报销流程,确保合规、及时、准确。", styles["BodyText"]))
    body.append(Paragraph("1.2 适用范围:全体正式员工,实习/外包员工参照执行。", styles["BodyText"]))

    body.append(Paragraph("<b>第二章 报销类别</b>", styles["Heading2"]))
    body.append(Paragraph("2.1 差旅交通:出差期间的机票/火车票/出租车。", styles["BodyText"]))
    body.append(Paragraph("2.2 业务招待:客户宴请、礼品。", styles["BodyText"]))
    body.append(Paragraph("2.3 办公采购:书籍、文具、低于 1000 元的办公用品。", styles["BodyText"]))
    body.append(Paragraph("2.4 培训费:经 HR 批准的外部培训。", styles["BodyText"]))

    body.append(Paragraph("<b>第三章 标准与限额</b>", styles["Heading2"]))
    body.append(Paragraph("3.1 国内差旅住宿:一线城市 800 元/晚,其他 600 元/晚。", styles["BodyText"]))
    body.append(Paragraph("3.2 业务招待:每人 300 元以内,需事前报批。", styles["BodyText"]))
    body.append(Paragraph("3.3 出租车票实报实销,需行程说明。", styles["BodyText"]))

    body.append(Paragraph("<b>第四章 流程</b>", styles["Heading2"]))
    body.append(Paragraph("4.1 员工 OA 填写报销单 → 上传发票 → 直属上级审批。", styles["BodyText"]))
    body.append(Paragraph("4.2 单笔 5000 元以下由部门负责人审批。", styles["BodyText"]))
    body.append(Paragraph("4.3 单笔 5000 元以上需总监 + 财务审批。", styles["BodyText"]))

    body.append(Paragraph("<b>第五章 票据要求</b>", styles["Heading2"]))
    body.append(Paragraph("5.1 必须提供合规增值税专用发票或普通发票。", styles["BodyText"]))
    body.append(Paragraph("5.2 电子发票需含发票号、开票日期、金额、销方税号。", styles["BodyText"]))

    body.append(Paragraph("<b>第六章 时效</b>", styles["Heading2"]))
    body.append(Paragraph("6.1 当月费用次月 10 日前提交,逾期需书面说明。", styles["BodyText"]))
    body.append(Paragraph("6.2 财务审核后 5 个工作日内付款。", styles["BodyText"]))

    doc.build(body)


def gen_09_tech_manual_pdf() -> None:
    """技术手册 — 大 PDF,165 标题结构(模拟真实技术文档体量)。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    doc = SimpleDocTemplate(str(OUT_DIR / "09_技术手册.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph(f"<b>{COMPANY} 跨境电商技术手册</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    body.append(Paragraph("<b>本手册涵盖跨境电商系统的架构、运维、安全规范。</b>", styles["BodyText"]))

    chapters = [
        "系统概述", "架构总览", "技术栈", "部署环境", "开发规范", "代码评审", "测试体系",
        "持续集成", "持续部署", "监控告警", "日志规范", "链路追踪", "性能基线",
        "容量评估", "限流策略", "熔断降级", "灾备方案", "数据库设计", "索引设计",
        "分库分表", "读写分离", "缓存策略", "消息队列", "任务调度", "分布式锁",
        "事务方案", "最终一致性", "幂等设计", "防重提交", "数据迁移", "数据归档",
        "数据备份", "数据恢复", "权限模型", "认证方案", "授权流程", "审计日志",
        "加密方案", "脱敏规范", "密码策略", "会话管理", "API 网关", "限流算法",
        "负载均衡", "服务发现", "配置中心", "灰度发布", "A/B 测试", "回滚预案",
        "故障响应", "值班机制", "升级窗口", "变更管理", "问题跟踪", "复盘流程",
        "知识库", "应急预案", "DDoS 防护", "WAF 规则", "爬虫识别", "风控引擎",
        "反欺诈", "实名认证", "支付链路", "对账系统", "资金安全", "合规审计",
        "个人信息保护", "跨境数据出境", "数据分级", "权限回收", "密钥轮换",
        "漏洞管理", "安全扫描", "渗透测试", "红蓝对抗", "威胁情报", "应急响应",
        "取证分析", "日志审计", "行为分析", "异常检测", "告警收敛", "根因分析",
        "故障演练", "混沌工程", "压测方案", "全链路追踪", "指标体系", "SLO 定义",
        "错误预算", "值班手册", "Runbook", "升级指南", "运维自动化", "配置管理",
        "发布管理", "版本控制", "Git 工作流", "Code Review 检查清单", "文档规范",
        "API 文档", "数据库文档", "架构决策记录", "复盘报告模板", "事故报告模板",
        "SOP 清单", "常见问题 FAQ", "联系支持", "版本历史", "贡献者名单", "许可证声明",
        "附 A:术语表", "附 B:缩略语", "附 C:参考资料", "附 D:变更记录",
        "附 E:FAQ 索引", "附 F:工具清单", "附 G:监控面板", "附 H:备份策略",
        "附 I:容灾切换", "附 J:应急联系人", "附 K:第三方依赖", "附 L:升级路径",
        "附 M:容量规划", "附 N:成本分析", "附 O:合规清单", "附 P:风险评估",
        "附 Q:法律声明", "附 R:商标说明", "附 S:版权声明", "附 T:致谢",
        "结语", "附录索引", "修订历史", "文档元数据", "反馈方式",
        "鸣谢",
    ]
    for i, ch in enumerate(chapters):
        body.append(Paragraph(f"<b>{i+1}. {ch}</b>", styles["Heading3"]))
        body.append(Paragraph(
            f"本节描述 {ch} 的设计原则与实现要点。系统采用主流架构,"
            f"关键指标包括可用性 ≥ 99.9%、P99 延迟 ≤ 500ms、错误率 ≤ 0.1%。"
            f"详细参数见对应章节表格与示例。",
            styles["BodyText"],
        ))
        body.append(Spacer(1, 6))
        if (i + 1) % 30 == 0:
            body.append(PageBreak())

    doc.build(body)


def gen_10_unstructured_pdf() -> None:
    """结构混乱的长文档 — PDF,大量短行/乱换行/重复,测 indexer 健壮性。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    doc = SimpleDocTemplate(str(OUT_DIR / "10_结构混乱的长文档.pdf"), pagesize=A4)
    styles = getSampleStyleSheet()
    apply_chinese_styles(styles, CHINESE_FONT)
    body = []
    body.append(Paragraph("<b>结构混乱文档</b>", styles["Title"]))
    body.append(Paragraph(DISCLAIMER, styles["Italic"]))
    body.append(Spacer(1, 12))

    # 模拟 OCR/扫描件导出 — 大段无标点的纯文本 + 重复
    fragments = [
        "本文件用于测试 RAG 系统的健壮性,内容由脚本随机生成,不代表任何实际业务流程。",
        "建议在导入系统前对原文做规范化处理,例如拆分段落、补全标点、统一大小写等。",
        "本段没有标题,读者难以快速定位主题,这也是结构混乱文档的常见问题。",
        "测试重点:Chunking 策略应当能识别段落边界并合理切分。",
        "",
        "重复声明:本文件仅用于内部测试,所有数据均为虚构,不构成任何业务承诺。",
        "重复声明:本文件仅用于内部测试,所有数据均为虚构,不构成任何业务承诺。",
        "重复声明:本文件仅用于内部测试,所有数据均为虚构,不构成任何业务承诺。",
        "",
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "",
        "无标点的连续中文片段本文件测试场景应该覆盖无标点纯文本,系统需要自行切分否则会出现单一大chunk导致检索失败",
        "",
        "1. 一、章节 (1) 列表混用",
        "2. 二、章节 (2) 列表混用",
        "3. 三、章节 (3) 列表混用",
        "4. 四、章节 (4) 列表混用",
        "5. 五、章节 (5) 列表混用",
        "",
        "中文与英文混排:本段包含example@example.com邮箱地址、HTTP://example.com链接、+86-138-0000-0000电话号码等敏感信息,系统应正确处理。",
        "",
        "TAB	分隔	的	文本	系统	应该	正确	解析	不要	因为	TAB	就	切错",
        "",
        "超长无段落标题示例标题A标题B标题C混在一起系统需要根据语义或标点拆开",
        "",
        "## Markdown 风格标题(应该是 Heading)",
        "内容:虽然 PDF 不识别 Markdown 语法,但 RAG 系统应该靠排版/字体识别这是标题。",
        "",
        "本页 PDF 字号 10.0 与正文一致,系统需要靠加粗识别层级。",
        "**加粗强调**(在 PDF 中表现为加粗,但 RAG 系统通常靠 fontSize + bold 识别)。",
        "",
        "数字清单: 1,000 1,000.00 1,000.0000 0.1 0.01 0.001 -100 +200 ± 0.5 100% 50‰",
        "日期: 2026-08-18 2026/08/18 2026年8月18日 2026.08.18 Aug 18, 2026",
        "金额: ¥1,000.00 $100 €50 £30 HK$200 NT$500",
        "",
        f"结束语:本文档结束于{DISCLAIMER}。",
    ]
    for f in fragments:
        if f == "":
            body.append(Spacer(1, 8))
        else:
            body.append(Paragraph(f, styles["BodyText"]))

    doc.build(body)


# ============ main ============

GENERATORS = [
    ("01_FAQ.md", gen_01_faq_md),
    ("02_售后FAQ.docx", gen_02_after_sales_faq_docx),
    ("03_库存管理制度.pdf", gen_03_inventory_pdf),
    ("04_采购流程.docx", gen_04_purchase_process_docx),
    ("05_商品规格说明.pdf", gen_05_product_spec_pdf),
    ("06_退货政策.pdf", gen_06_return_policy_pdf),
    ("07_物流SOP.docx", gen_07_logistics_sop_docx),
    ("08_员工报销制度.pdf", gen_08_reimbursement_pdf),
    ("09_技术手册.pdf", gen_09_tech_manual_pdf),
    ("10_结构混乱的长文档.pdf", gen_10_unstructured_pdf),
    ("11_采购合同.md", gen_11_purchase_contract_md),
]


def main() -> int:
    print(f"输出目录: {OUT_DIR}")
    for name, fn in GENERATORS:
        try:
            fn()
            size = (OUT_DIR / name).stat().st_size
            print(f"  [OK] {name}  ({size:,} bytes)")
        except Exception as e:
            print(f"  [FAIL] {name}: {e}")
            return 1
    print(f"\n生成完成: {len(GENERATORS)} 个文件")
    return 0


if __name__ == "__main__":
    sys.exit(main())