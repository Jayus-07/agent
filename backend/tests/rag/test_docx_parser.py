"""test_docx_parser.py — DocxParser 单元测试。

覆盖：
1. 基本加载：Heading 1/2 + 段落识别
2. 表格识别：表格 → table 节点，rows 不为空
3. 章节归属：paragraph 挂在最近 section 下
4. Heading 层级 pop 逻辑：Heading 1 → Heading 3 → Heading 1 正确嵌套
"""
import docx
import pytest

from backend.rag.preprocessing.ast import walk
from backend.rag.preprocessing.parser.docx_parser import DocxParser


@pytest.fixture
def sample_docx(tmp_path):
    """创建 docx：1 个 Heading 1 + 1 个 Heading 2 + 2 段 + 1 表。"""
    p = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_heading("一级标题", level=1)
    d.add_paragraph("一级标题下的段落。")
    d.add_heading("二级标题", level=2)
    d.add_paragraph("二级标题下的段落。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "列2"
    table.cell(1, 0).text = "数据1"
    table.cell(1, 1).text = "数据2"
    d.save(str(p))
    return str(p)


def test_docx_parser_basic_load(sample_docx):
    ast = DocxParser().parse(sample_docx)
    assert ast.source_file == sample_docx
    assert ast.raw_text != ""
    sections = [
        (n.text, n.level) for n in walk(ast.root)
        if n.type == "section" and n.level > 0
    ]
    assert ("一级标题", 1) in sections
    assert ("二级标题", 2) in sections


def test_docx_parser_table_recognized(sample_docx):
    ast = DocxParser().parse(sample_docx)
    tables = [n for n in walk(ast.root) if n.type == "table"]
    assert len(tables) == 1
    assert tables[0].rows is not None
    assert len(tables[0].rows) == 2


def test_docx_parser_section_attribution(sample_docx):
    """paragraph 应挂在最近的 section 下。"""
    ast = DocxParser().parse(sample_docx)
    h1_section = next(
        n for n in walk(ast.root)
        if n.type == "section" and n.text == "一级标题"
    )
    h1_paragraphs = [c for c in h1_section.children if c.type == "paragraph"]
    assert any("一级标题下的段落" in c.text for c in h1_paragraphs)


def test_docx_parser_heading_level_pop_logic(tmp_path):
    """Heading 3 后出现 Heading 1 → Heading 1 应挂 root，Heading 3 保留嵌套。"""
    p = tmp_path / "nested.docx"
    d = docx.Document()
    d.add_heading("外层 1", level=1)
    d.add_heading("内层 3", level=3)
    d.add_paragraph("内层内容。")
    d.add_heading("回到外层 1", level=1)  # 回到 level=1
    d.add_paragraph("外层内容。")
    d.save(str(p))

    ast = DocxParser().parse(str(p))
    sections = [
        (n.text, n.level) for n in walk(ast.root)
        if n.type == "section" and n.level > 0
    ]
    # 两个 level=1 section 都在 root 下，level=3 在第一个 level=1 下
    assert ("外层 1", 1) in sections
    assert ("回到外层 1", 1) in sections
    assert ("内层 3", 3) in sections

    # 验证"内层 3" 挂在第一个 "外层 1" 下，不是 root
    outer1 = next(
        n for n in walk(ast.root)
        if n.type == "section" and n.text == "外层 1"
    )
    inner3_in_outer1 = any(
        c.type == "section" and c.text == "内层 3"
        for c in outer1.children
    )
    assert inner3_in_outer1


def test_docx_parser_table_text_preserved(sample_docx):
    """表格内容应进入 table 节点的 text，供 chunk 保留（避免空 chunk 丢失表格数据）。"""
    ast = DocxParser().parse(sample_docx)
    tables = [n for n in walk(ast.root) if n.type == "table"]
    assert len(tables) == 1
    # table.text 非空，且包含单元格内容
    table_text = tables[0].text
    assert "列1" in table_text
    assert "数据2" in table_text


def test_docx_parser_qa_doc(tmp_path):
    """DOCX FAQ（Q：/A： 格式）应识别为 QA 文档，产 qa_question/qa_answer 节点。"""
    p = tmp_path / "faq.docx"
    d = docx.Document()
    d.add_paragraph("Q：签收后发现商品破损怎么办？")
    d.add_paragraph("A：请在签收后 48 小时内拍照留证。")
    d.add_paragraph("Q：换货的流程是怎样的？")
    d.add_paragraph("A：提交换货申请。")
    d.save(str(p))

    ast = DocxParser().parse(str(p))
    qa_nodes = [
        n for n in ast.root.children if n.type in ("qa_question", "qa_answer")
    ]
    assert len(qa_nodes) == 4  # 2 问题 + 2 答案
    assert qa_nodes[0].type == "qa_question"
    assert "破损" in qa_nodes[0].text
    assert qa_nodes[1].type == "qa_answer"
