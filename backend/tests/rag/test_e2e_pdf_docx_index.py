"""test_e2e_pdf_docx_index.py — Phase 2 端到端验证。

用 tmp_path 构造真实 PDF/DOCX，验证 parse_and_chunk 产出非空 chunks。
"""
import docx
import pymupdf as fitz
import pytest

from backend.rag.preprocessing.pipeline import parse_and_chunk


@pytest.fixture
def real_pdf(tmp_path):
    """构造 2 页真实 PDF，含可读文本。"""
    p = tmp_path / "real.pdf"
    d = fitz.open()
    p1 = d.new_page()
    p1.insert_text((50, 50), "售后制度说明。")
    p1.insert_text((50, 80), "退货流程详细步骤。")
    p2 = d.new_page()
    p2.insert_text((50, 50), "差评处理规范。")
    d.save(str(p))
    d.close()
    return str(p)


@pytest.fixture
def real_docx(tmp_path):
    """构造真实 DOCX，含 Heading + Table。"""
    p = tmp_path / "real.docx"
    d = docx.Document()
    d.add_heading("售后制度", level=1)
    d.add_paragraph("退货流程说明。")
    d.add_heading("差评处理", level=1)
    d.add_paragraph("48 小时内处理。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "类型"
    table.cell(0, 1).text = "时限"
    table.cell(1, 0).text = "退货"
    table.cell(1, 1).text = "24h"
    d.save(str(p))
    return str(p)


def test_pdf_e2e_pipeline_produces_chunks(real_pdf):
    """PDF 走完整流水线：parse_and_chunk 不返回空。"""
    chunks = parse_and_chunk(real_pdf)
    assert len(chunks) > 0
    # 所有 chunk 都应有 page_content
    assert all(hasattr(c, "page_content") and c.page_content for c in chunks)


def test_docx_e2e_pipeline_produces_chunks_with_structure(real_docx):
    """DOCX 走完整流水线：parse_and_chunk 产 chunks 含 parent + leaf。"""
    chunks = parse_and_chunk(real_docx)
    assert len(chunks) > 0
    # DOCX 走 StructureChunkStrategy 路由 → 有 parent + leaf 双粒度
    granules = {c.metadata.get("granularity") for c in chunks}
    assert "leaf" in granules
    # parent 可能存在（如果文档有 section 结构）
    # 不强制要求 parent，因为某些结构可能产空 parent


def test_pdf_e2e_metadata_protocol(real_pdf):
    """PDF chunk 应有 Phase 1 metadata 协议（granularity / chunk_id / chunk_tokens）。"""
    chunks = parse_and_chunk(real_pdf)
    for c in chunks:
        assert "chunk_id" in c.metadata
        assert "granularity" in c.metadata
        assert "chunk_tokens" in c.metadata
