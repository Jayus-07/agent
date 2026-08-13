"""DocxParser — python-docx 解析 Word → Raw AST。

段落分类规则：
- style.name 匹配 `Heading {N}` → DocumentNode(type="section", level=N, text=title)
- 段落 style.name 含 `List` → DocumentNode(type="list")
- 表格 → DocumentNode(type="table", rows=...)
- 其他 → DocumentNode(type="paragraph")

章节归属：paragraph / list / table 挂在最近的祖先 section 下。
"""
from __future__ import annotations

import re

import docx

from backend.rag.preprocessing.ast import DocumentAST, DocumentNode
from backend.rag.preprocessing.parser._qa_patterns import (
    extract_qa_pairs, looks_like_qa_doc,
)
from backend.rag.preprocessing.parser.base import BaseDocumentParser
from backend.shared.logger import logger

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
_LIST_RE = re.compile(r"List\s+\w+", re.IGNORECASE)


def _find_parent_section(
    stack: list[DocumentNode], level: int
) -> DocumentNode:
    """在 section 栈中找到新 section 应该挂的父节点。

    规则：弹出所有 level >= 当前 level 的祖先，让新 section 挂在
    第一个 level < 当前 level 的祖先下。栈底始终是 root（level=0）。
    """
    while len(stack) > 1 and stack[-1].level >= level:
        stack.pop()
    return stack[-1]


class DocxParser(BaseDocumentParser):
    def parse(self, file_path: str) -> DocumentAST:
        try:
            d = docx.Document(file_path)
        except (docx.opc.exceptions.PackageNotFoundError, ValueError) as e:
            # 文件损坏 / 格式异常 → 返回空 AST
            logger.error(
                f"[DocxParser] 打开失败 {file_path}: "
                f"{type(e).__name__}: {e}"
            )
            return DocumentAST(
                root=DocumentNode(type="section", text="", level=0),
                source_file=file_path,
                raw_text="",
            )
        except Exception as e:
            logger.exception(
                f"[DocxParser] 打开异常 {file_path}: {type(e).__name__}"
            )
            raise

        root = DocumentNode(type="section", text="", level=0)

        # QA 识别：拼接段落文本，判断是否 FAQ 文档（Q：/A： 等格式）。
        # 命中则整篇按 Q/A 切（产 qa_question/qa_answer 节点），不走普通段落分类，
        # 与 MarkdownParser 的 FAQ 路径保持一致。
        full_text = "\n".join(
            p.text.strip() for p in d.paragraphs if p.text.strip()
        )
        if looks_like_qa_doc(full_text):
            pairs = extract_qa_pairs(full_text)
            for q, a, _ptype in pairs:
                root.children.append(DocumentNode(type="qa_question", text=q))
                root.children.append(DocumentNode(type="qa_answer", text=a))
            logger.info(
                f"[DocxParser] {file_path} 识别为 FAQ 文档，"
                f"产出 {len(pairs)} 个 Q/A 对"
            )
            return DocumentAST(root=root, source_file=file_path, raw_text=full_text)

        section_stack: list[DocumentNode] = [root]
        raw_lines: list[str] = []
        skipped_tables: list[int] = []

        # 段落
        for para in d.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            heading_match = _HEADING_RE.match(style_name)
            if heading_match and text:
                level = int(heading_match.group(1))
                section = DocumentNode(type="section", text=text, level=level)
                parent = _find_parent_section(section_stack, level)
                parent.children.append(section)
                section_stack.append(section)
                raw_lines.append(text)
                continue

            if _LIST_RE.search(style_name) and text:
                list_node = DocumentNode(type="list", text=text)
                section_stack[-1].children.append(list_node)
                raw_lines.append(text)
                continue

            if text:
                para_node = DocumentNode(type="paragraph", text=text)
                section_stack[-1].children.append(para_node)
                raw_lines.append(text)

        # 表格（python-docx 表格不在段落流中，需单独遍历）
        for tbl_idx, table in enumerate(d.tables):
            try:
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]
            except (AttributeError, IndexError) as e:
                logger.warning(
                    f"[DocxParser] 表格 {tbl_idx} 解析失败: "
                    f"{type(e).__name__}: {e}"
                )
                skipped_tables.append(tbl_idx)
                continue
            # 表格内容同时写入 text 字段：否则 StructureChunkStrategy 用 text 切分
            # 会产出空 leaf chunk（ChunkFilter 拒绝），表格数据入库后彻底丢失
            table_text = "\n".join(", ".join(r) for r in rows)
            table_node = DocumentNode(type="table", text=table_text, rows=rows)
            section_stack[-1].children.append(table_node)
            raw_lines.append(table_text)

        # 可观测：汇总报告
        if skipped_tables:
            logger.warning(
                f"[DocxParser] {file_path} 跳过 {len(skipped_tables)} 个表: "
                f"{skipped_tables}"
            )

        raw_text = "\n".join(raw_lines)
        return DocumentAST(root=root, source_file=file_path, raw_text=raw_text)
