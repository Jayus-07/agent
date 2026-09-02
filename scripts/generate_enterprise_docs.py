#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成企业文档（.docx）— 虚构跨境电商知识库补充。
使用 D:\Python (3.10.2) 中的 python-docx 生成，与项目托管 Python 隔离。
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"D:\Program Files\workplace\agent\data\docs\rag_test_kb\general"


def set_cjk(run, name="宋体", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cjk(r, "黑体", 16)
    r.bold = True


def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cjk(r, "黑体", 13)
    r.bold = True


def p(doc, text, size=11):
    para = doc.add_paragraph()
    set_cjk(para.add_run(text), "宋体", size)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    set_cjk(para.add_run(text), "宋体", 11)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(head)
        set_cjk(r, "黑体", 10)
        r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(c))
            set_cjk(r, "宋体", 10)


def build_qc():
    doc = Document()
    title(doc, "入库质检标准作业程序（SOP）")
    p(doc, "【仅供测试 · 内容虚构】", size=10)
    h(doc, "一、目的")
    p(doc, "规范跨境商品入库质检流程，确保上架商品符合质量与合规要求，降低售后退货与客诉风险。")
    h(doc, "二、适用范围")
    p(doc, "适用于示例跨境电商公司所有自营及代发商品的入库质检，含境内仓与海外仓。")
    h(doc, "三、抽检比例（AQL）")
    table(doc, ["批量区间（件）", "抽检比例", "最少抽检数"],
          [["≤500", "5%", "20"], ["501–2000", "3%", "25"], [">2000", "2%", "40"]])
    h(doc, "四、检验项目")
    bullet(doc, "外观：划痕、变形、色差、污渍")
    bullet(doc, "功能：通电/试用核心功能正常")
    bullet(doc, "标签：含 CE/FCC 标识、HS 编码、中文说明（如适用）")
    bullet(doc, "配件：说明书、保修卡、适配器齐全")
    h(doc, "五、合格判定")
    p(doc, "批次不合格率 ≤5% 整批通过；>5% 整批拒收或要求换货。")
    h(doc, "六、不合格品处理")
    bullet(doc, "隔离：移至不合格品区并贴红标")
    bullet(doc, "通知：30 分钟内通知采购部")
    bullet(doc, "响应：供应商须在 48 小时内响应")
    bullet(doc, "处置：退/换/返工，返工后按原比例复检")
    h(doc, "七、记录与归档")
    p(doc, "质检单（含抽检数、不合格项、处置结果）归档保存 2 年，供追溯与供应商考核使用。")
    out = BASE + r"\13_质检标准SOP.docx"
    doc.save(out)
    print("saved:", out)


def build_hr():
    doc = Document()
    title(doc, "员工手册（HR 制度）")
    p(doc, "【仅供测试 · 内容虚构】", size=10)
    h(doc, "一、入职与转正")
    p(doc, "试用期 3 个月；转正须直属主管评估 + HR 审批，转正后享受全额福利。")
    h(doc, "二、工作时间与考勤")
    p(doc, "标准工时 9:30–18:30，弹性 ±30 分钟；月度缺勤（事假+病假）累计 ≤3 天为正常。")
    h(doc, "三、假期")
    table(doc, ["假别", "天数/规则", "说明"],
          [["年假", "5–15 天（按司龄）", "司龄满 1 年起算"],
           ["病假", "凭医院证明", "三甲或社区医院诊断"],
           ["事假", "提前 1 天申请", "月度 ≤2 天"],
           ["婚假/产假", "依法执行", "提供相应证明"]])
    h(doc, "四、薪酬与绩效")
    p(doc, "每月 10 日发薪；季度绩效奖金按考核结果发放；年终奖依公司利润水平而定。")
    h(doc, "五、行为准则")
    bullet(doc, "禁止收受供应商礼品折现价值 >200 元，超额须登记上交")
    bullet(doc, "保守商业秘密，保密义务期限见劳动合同与采购合同相关条款")
    bullet(doc, "利益冲突须主动申报")
    h(doc, "六、奖惩")
    p(doc, "设立年度突出贡献奖；违规按《员工奖惩办法》处理，严重者解除劳动合同。")
    h(doc, "七、离职")
    p(doc, "离职须提前 30 日书面申请，完成工作交接清单签字后方可结算薪资。")
    out = BASE + r"\15_员工手册HR制度.docx"
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    build_qc()
    build_hr()
    print("DONE")
